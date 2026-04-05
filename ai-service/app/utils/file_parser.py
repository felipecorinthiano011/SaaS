"""Utility functions for parsing resume files."""

import json
import re
from typing import Optional
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None


def parse_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    import logging
    logger = logging.getLogger(__name__)

    if not pdfplumber:
        raise ImportError("pdfplumber is not installed. Install with: pip install pdfplumber")

    text = []
    try:
        logger.info(f"Opening PDF file: {file_path}")
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} pages")
            for idx, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    logger.info(f"Page {idx+1}: extracted {len(page_text)} characters")
                    text.append(page_text)
                else:
                    logger.warning(f"Page {idx+1}: no text extracted (might be image-only)")

        result = "\n".join(text)
        logger.info(f"Total extracted from PDF: {len(result)} characters")
        return result
    except Exception as e:
        logger.error(f"Failed to parse PDF: {str(e)}", exc_info=True)
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    import logging
    logger = logging.getLogger(__name__)

    if not Document:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")

    text = []
    try:
        logger.info(f"Opening DOCX file: {file_path}")
        doc = Document(file_path)
        logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")

        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text.append(paragraph.text)

        result = "\n".join(text)
        logger.info(f"Total extracted from DOCX: {len(result)} characters")
        return result
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {str(e)}", exc_info=True)
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_resume_file(file_path: str, file_name: str) -> str:
    """Parse resume file (PDF or DOCX) and extract text."""
    file_name_lower = file_name.lower()

    if file_name_lower.endswith('.pdf'):
        return parse_pdf(file_path)
    elif file_name_lower.endswith('.docx') or file_name_lower.endswith('.doc'):
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_name}. Only PDF and DOCX are supported.")


def extract_json_from_text(text: str) -> dict:
    """Extract JSON object from text response."""
    # Find first { and last } in the text
    start_idx = text.find('{')
    end_idx = text.rfind('}') + 1

    if start_idx == -1 or end_idx == 0:
        raise ValueError("No JSON object found in response")

    json_str = text[start_idx:end_idx]
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse JSON: {str(e)}")


def clean_text(text: str) -> str:
    """Clean and normalize text for AI processing.

    Removes:
    - Email addresses
    - URLs and links
    - Phone numbers
    - Extra whitespace
    - Headers/footers with URLs

    Args:
        text: Raw text to clean

    Returns:
        Clean plain text suitable for AI processing
    """
    # Remove email addresses
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '', text)

    # Remove URLs and links (http, https, www, ftp)
    text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    text = re.sub(r'www\.[a-zA-Z0-9-]+\.[a-zA-Z]{2,}', '', text)
    text = re.sub(r'ftp://[^\s]+', '', text)

    # Remove LinkedIn URLs and similar social profiles
    text = re.sub(r'linkedin\.com/in/[^\s]+', '', text)
    text = re.sub(r'github\.com/[^\s]+', '', text)
    text = re.sub(r'twitter\.com/[^\s]+', '', text)

    # Remove phone numbers (various formats)
    text = re.sub(r'\b(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b', '', text)

    # Remove common header/footer patterns
    # Remove page numbers (e.g., "Page 1", "- 1 -", etc.)
    text = re.sub(r'[-\s]*page\s+\d+[-\s]*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'[-\s]*\d+\s+of\s+\d+[-\s]*', '', text, flags=re.IGNORECASE)

    # Remove lines that are just dashes, equals signs, or asterisks (common separators)
    text = re.sub(r'^[\s]*[-_=*]{3,}[\s]*$', '', text, flags=re.MULTILINE)

    # Remove multiple consecutive spaces and special separator characters
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'([-=_*]){3,}', '', text)

    # Remove multiple consecutive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Strip whitespace from each line and remove empty lines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    text = '\n'.join(lines)

    # Final trim
    text = text.strip()

    return text


def extract_resume_text(file) -> str:
    """Extract and clean text from resume file (PDF or DOCX).

    This function:
    1. Detects file format (PDF or DOCX)
    2. Extracts raw text from the file
    3. Cleans the text (removes emails, links, phone numbers, extra whitespace)
    4. Returns plain text suitable for AI processing

    Args:
        file: File-like object or file path (str or Path)

    Returns:
        Clean plain text extracted from the resume

    Raises:
        ValueError: If file format is not supported or parsing fails
        ImportError: If required library (pdfplumber or python-docx) is not installed

    Example:
        >>> text = extract_resume_text(resume_file)
        >>> print(text)  # Clean resume text without emails/links
    """
    import logging
    logger = logging.getLogger(__name__)

    # Convert file path to string if Path object
    if isinstance(file, Path):
        file_path = str(file)
    elif isinstance(file, str):
        file_path = file
    else:
        # Assume it's a file-like object with a name attribute
        if hasattr(file, 'filename'):
            file_name = file.filename
            logger.info(f"Processing uploaded file: {file_name}")

            # For file objects, we need to save to temporary location
            import tempfile
            import shutil

            temp_file = tempfile.NamedTemporaryFile(
                suffix=Path(file_name).suffix,
                delete=False
            )
            try:
                shutil.copyfileobj(file.file, temp_file)
                temp_file.close()
                file_path = temp_file.name
                file_name_lower = file_name.lower()

                logger.info(f"File size: {Path(file_path).stat().st_size} bytes")

                # Parse the file
                if file_name_lower.endswith('.pdf'):
                    logger.info("Parsing PDF file...")
                    raw_text = parse_pdf(file_path)
                elif file_name_lower.endswith(('.docx', '.doc')):
                    logger.info("Parsing DOCX file...")
                    raw_text = parse_docx(file_path)
                else:
                    raise ValueError(f"Unsupported file format: {file_name}. Only PDF and DOCX are supported.")

                logger.info(f"Raw text extracted: {len(raw_text)} characters")

                # Clean and return
                cleaned_text = clean_text(raw_text)
                logger.info(f"After cleaning: {len(cleaned_text)} characters")

                if not cleaned_text.strip():
                    logger.warning("Warning: Extracted text is empty after cleaning")

                return cleaned_text
            finally:
                # Clean up temporary file
                import os
                if os.path.exists(file_path):
                    os.unlink(file_path)
                    logger.info("Temporary file deleted")
        else:
            raise ValueError("File input must be a file path (str/Path) or a file-like object with filename")

    # Parse the file based on extension
    raw_text = parse_resume_file(file_path, file_path)

    # Clean and return
    return clean_text(raw_text)






